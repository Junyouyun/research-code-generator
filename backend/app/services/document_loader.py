import re
from pathlib import Path

from app.services.document_elements import make_document_element


SUPPORTED_EXTENSIONS = {".pdf", ".md", ".markdown", ".txt", ".docx"}


def load_document_elements(file_path: Path, document_id: str, original_filename: str = "") -> list[dict]:
    suffix = _get_suffix(file_path, original_filename)
    if suffix not in SUPPORTED_EXTENSIONS:
        raise RuntimeError(f"不支持的文档类型：{suffix or 'unknown'}")

    if suffix == ".pdf":
        return _load_pdf(file_path, document_id)
    if suffix in {".md", ".markdown"}:
        return _load_markdown(file_path, document_id, suffix.lstrip("."))
    if suffix == ".txt":
        return _load_text(file_path, document_id)
    if suffix == ".docx":
        return _load_docx(file_path, document_id)

    raise RuntimeError(f"不支持的文档类型：{suffix}")


def elements_to_parsed_paper(elements: list[dict], source: str) -> dict:
    pages = _elements_to_pages(elements)
    full_text = "\n\n".join(element["text"] for element in elements if element.get("text"))

    return {
        "source": source,
        "title": _guess_title(elements),
        "abstract": _guess_abstract(elements),
        "page_count": len(pages) or None,
        "pages": pages or [{"page_number": 1, "text": full_text}],
        "sections": _elements_to_sections(elements),
        "full_text": full_text,
        "tables": [element for element in elements if element.get("type") == "table"],
        "figures": [element for element in elements if element.get("type") == "figure"],
        "references": [element for element in elements if element.get("type") == "reference"],
    }


def _load_pdf(file_path: Path, document_id: str) -> list[dict]:
    try:
        import fitz
    except ModuleNotFoundError as exc:
        raise RuntimeError("缺少 PyMuPDF，请先安装 backend/requirements.txt 中的依赖。") from exc

    elements = []
    order_index = 1
    current_section = ""

    with fitz.open(file_path) as document:
        repeated_texts = _find_repeated_pdf_margin_texts(document)
        for page_index, page in enumerate(document, start=1):
            blocks = _sort_pdf_blocks(page.get_text("blocks"), page.rect.width)
            for block in blocks:
                block_type = block[6] if len(block) > 6 else 0
                text = _clean_block_text(block[4] if len(block) > 4 else "")

                if block_type == 1:
                    elements.append(
                        make_document_element(
                            document_id=document_id,
                            source_file_type="pdf",
                            element_type="figure",
                            text=f"[Image] page {page_index}",
                            markdown=f"[Image] page {page_index}",
                            order_index=order_index,
                            page_start=page_index,
                            page_end=page_index,
                            bbox=[float(block[0]), float(block[1]), float(block[2]), float(block[3])],
                            section_title=current_section,
                            hierarchy_path=current_section,
                            confidence=0.5,
                            needs_review=True,
                            extra={"source": "pdf_image_block"},
                        )
                    )
                    order_index += 1
                    continue

                if not text:
                    continue
                if text in repeated_texts:
                    continue

                element_type = _guess_text_type(text)
                if element_type == "title":
                    current_section = text.splitlines()[0]

                if _is_reference_section(current_section) and element_type != "title":
                    for reference in _split_reference_entries(text):
                        elements.append(
                            make_document_element(
                                document_id=document_id,
                                source_file_type="pdf",
                                element_type="reference",
                                text=reference,
                                order_index=order_index,
                                page_start=page_index,
                                page_end=page_index,
                                bbox=[float(block[0]), float(block[1]), float(block[2]), float(block[3])],
                                section_title=current_section,
                                hierarchy_path=current_section,
                                confidence=0.7,
                            )
                        )
                        order_index += 1
                    continue

                elements.append(
                    make_document_element(
                        document_id=document_id,
                        source_file_type="pdf",
                        element_type=element_type,
                        text=text,
                        order_index=order_index,
                        page_start=page_index,
                        page_end=page_index,
                        bbox=[float(block[0]), float(block[1]), float(block[2]), float(block[3])],
                        section_title=current_section,
                        hierarchy_path=current_section,
                        confidence=0.8 if element_type == "title" else 0.7,
                    )
                )
                order_index += 1

    return elements


def _load_markdown(file_path: Path, document_id: str, source_file_type: str) -> list[dict]:
    content = file_path.read_text(encoding="utf-8")
    elements = []
    order_index = 1
    heading_stack: list[str] = []
    paragraph_lines: list[str] = []
    special_lines: list[str] = []
    special_type = ""

    def add_element(element_type: str, lines: list[str], markdown: str | None = None) -> None:
        nonlocal order_index
        text = "\n".join(lines).strip()
        if not text:
            return
        elements.append(
            make_document_element(
                document_id=document_id,
                source_file_type=source_file_type,
                element_type=element_type,
                text=_strip_markdown_heading(text) if element_type == "title" else text,
                markdown=markdown or text,
                order_index=order_index,
                section_title=heading_stack[-1] if heading_stack else "",
                hierarchy_path=" > ".join(heading_stack),
                confidence=1.0,
            )
        )
        order_index += 1

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if not paragraph_lines:
            return
        text = "\n".join(paragraph_lines).strip()
        add_element("formula" if _looks_like_formula(text) else "paragraph", paragraph_lines)
        paragraph_lines = []

    def flush_special() -> None:
        nonlocal special_lines, special_type
        add_element(special_type, special_lines)
        special_lines = []
        special_type = ""

    for line in content.splitlines():
        if special_type == "code":
            special_lines.append(line)
            if line.strip().startswith("```"):
                flush_special()
            continue

        if special_type == "formula":
            special_lines.append(line)
            if line.strip().endswith("$$") and len(special_lines) > 1:
                flush_special()
            continue

        if line.strip().startswith("```"):
            flush_paragraph()
            special_type = "code"
            special_lines = [line]
            continue

        if line.strip().startswith("$$"):
            flush_paragraph()
            special_type = "formula"
            special_lines = [line]
            if line.strip().endswith("$$") and len(line.strip()) > 2:
                flush_special()
            continue

        if _is_markdown_image(line):
            flush_paragraph()
            add_element("figure", [line])
            continue

        if _is_markdown_heading(line):
            flush_paragraph()
            level, title = _parse_markdown_heading(line)
            heading_stack = heading_stack[: level - 1] + [title]
            add_element("title", [title], markdown=line)
            continue

        if _is_markdown_table_line(line):
            flush_paragraph()
            if special_type and special_type != "table":
                flush_special()
            special_type = "table"
            special_lines.append(line)
            continue

        if special_type == "table":
            flush_special()

        if _is_list_line(line):
            flush_paragraph()
            add_element("list", [line])
            continue

        if not line.strip():
            flush_paragraph()
            continue

        paragraph_lines.append(line)

    flush_paragraph()
    if special_type:
        flush_special()

    return elements


def _load_text(file_path: Path, document_id: str) -> list[dict]:
    content = file_path.read_text(encoding="utf-8")
    elements = []
    current_section = ""
    order_index = 1
    in_references = False

    for paragraph in _split_paragraphs(content):
        element_type = _guess_text_type(paragraph)
        if element_type == "title":
            current_section = paragraph.splitlines()[0]
            in_references = _is_reference_section(current_section)

        if in_references and element_type != "title":
            for reference in _split_reference_entries(paragraph):
                elements.append(
                    make_document_element(
                        document_id=document_id,
                        source_file_type="txt",
                        element_type="reference",
                        text=reference,
                        order_index=order_index,
                        section_title=current_section,
                        hierarchy_path=current_section,
                        confidence=0.7,
                    )
                )
                order_index += 1
            continue

        elements.append(
            make_document_element(
                document_id=document_id,
                source_file_type="txt",
                element_type=element_type,
                text=paragraph,
                order_index=order_index,
                section_title=current_section,
                hierarchy_path=current_section,
                confidence=0.7 if element_type == "title" else 0.6,
            )
        )
        order_index += 1
    return elements


def _load_docx(file_path: Path, document_id: str) -> list[dict]:
    try:
        from docx import Document
    except ModuleNotFoundError as exc:
        raise RuntimeError("缺少 python-docx，请先安装 backend/requirements.txt 中的依赖。") from exc

    document = Document(file_path)
    elements = []
    order_index = 1
    heading_stack: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
        element_type = "title" if style_name.startswith("heading") else _guess_text_type(text)
        if element_type == "title":
            level = _heading_level(style_name)
            heading_stack = heading_stack[: level - 1] + [text]

        elements.append(
            make_document_element(
                document_id=document_id,
                source_file_type="docx",
                element_type=element_type,
                text=text,
                order_index=order_index,
                section_title=heading_stack[-1] if heading_stack else "",
                hierarchy_path=" > ".join(heading_stack),
                confidence=0.9 if element_type == "title" else 0.8,
            )
        )
        order_index += 1

    for table in document.tables:
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        elements.append(
            make_document_element(
                document_id=document_id,
                source_file_type="docx",
                element_type="table",
                text="\n".join(rows),
                markdown="\n".join(rows),
                order_index=order_index,
                section_title=heading_stack[-1] if heading_stack else "",
                hierarchy_path=" > ".join(heading_stack),
                confidence=0.8,
            )
        )
        order_index += 1

    return elements


def _load_docx(file_path: Path, document_id: str) -> list[dict]:
    try:
        from docx import Document
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ModuleNotFoundError as exc:
        raise RuntimeError("缺少 python-docx，请先安装 backend/requirements.txt 中的依赖。") from exc

    document = Document(file_path)
    elements = []
    order_index = 1
    heading_stack: list[str] = []
    in_references = False

    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, document)
            text = paragraph.text.strip()
            if not text:
                continue

            style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
            element_type = "title" if style_name.startswith("heading") else _guess_text_type(text)
            if element_type == "title":
                level = _heading_level(style_name)
                heading_stack = heading_stack[: level - 1] + [text]
                in_references = _is_reference_section(text)

            if in_references and element_type != "title":
                element_type = "reference"

            elements.append(
                make_document_element(
                    document_id=document_id,
                    source_file_type="docx",
                    element_type=element_type,
                    text=text,
                    order_index=order_index,
                    section_title=heading_stack[-1] if heading_stack else "",
                    hierarchy_path=" > ".join(heading_stack),
                    confidence=0.9 if element_type == "title" else 0.8,
                )
            )
            order_index += 1
        elif isinstance(child, CT_Tbl):
            table = Table(child, document)
            rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
            elements.append(
                make_document_element(
                    document_id=document_id,
                    source_file_type="docx",
                    element_type="table",
                    text="\n".join(rows),
                    markdown="\n".join(rows),
                    order_index=order_index,
                    section_title=heading_stack[-1] if heading_stack else "",
                    hierarchy_path=" > ".join(heading_stack),
                    confidence=0.8,
                )
            )
            order_index += 1

    return elements


def _get_suffix(file_path: Path, original_filename: str) -> str:
    source = Path(original_filename).suffix if original_filename else file_path.suffix
    return source.lower()


def _guess_text_type(text: str) -> str:
    value = " ".join(text.split())
    if not value:
        return "unknown"
    if _looks_like_reference(value):
        return "reference"
    if _looks_like_heading(value):
        return "title"
    if _looks_like_table(value):
        return "table"
    if _looks_like_formula(value):
        return "formula"
    return "paragraph"


def _looks_like_heading(value: str) -> bool:
    if len(value) > 120:
        return False
    lowered = value.lower().strip()
    heading_words = {
        "abstract",
        "keywords",
        "introduction",
        "related work",
        "background",
        "method",
        "methods",
        "methodology",
        "approach",
        "experiment",
        "experiments",
        "evaluation",
        "results",
        "discussion",
        "conclusion",
        "references",
        "摘要",
        "关键词",
        "引言",
        "方法",
        "实验",
        "结果",
        "结论",
        "参考文献",
    }
    if lowered in heading_words:
        return True
    return bool(re.match(r"^\d+(?:\.\d+)*\.?\s+\S+", value))


def _looks_like_reference(value: str) -> bool:
    return bool(re.match(r"^\[\d+\]\s+", value) or re.match(r"^\d+\.\s+\S+", value))


def _sort_pdf_blocks(blocks: list[tuple], page_width: float) -> list[tuple]:
    text_blocks = [block for block in blocks if (block[6] if len(block) > 6 else 0) != 1]
    image_blocks = [block for block in blocks if (block[6] if len(block) > 6 else 0) == 1]

    if not _looks_like_two_columns(text_blocks, page_width):
        return sorted(blocks, key=lambda block: (round(block[1], 1), round(block[0], 1)))

    midpoint = page_width / 2
    left = [block for block in text_blocks if block[0] < midpoint]
    right = [block for block in text_blocks if block[0] >= midpoint]
    sorted_text = sorted(left, key=lambda block: (round(block[1], 1), round(block[0], 1)))
    sorted_text += sorted(right, key=lambda block: (round(block[1], 1), round(block[0], 1)))
    return sorted_text + sorted(image_blocks, key=lambda block: (round(block[1], 1), round(block[0], 1)))


def _looks_like_two_columns(blocks: list[tuple], page_width: float) -> bool:
    if len(blocks) < 6:
        return False
    left_count = sum(1 for block in blocks if block[0] < page_width * 0.45)
    right_count = sum(1 for block in blocks if block[0] > page_width * 0.45)
    return left_count >= 3 and right_count >= 3


def _find_repeated_pdf_margin_texts(document) -> set[str]:
    candidates: dict[str, int] = {}
    if len(document) < 3:
        return set()

    for page in document:
        height = page.rect.height
        for block in page.get_text("blocks"):
            text = _clean_block_text(block[4] if len(block) > 4 else "")
            if not text:
                continue
            in_margin = block[1] < height * 0.08 or block[3] > height * 0.92
            if in_margin and len(text) <= 120:
                candidates[text] = candidates.get(text, 0) + 1

    threshold = max(2, len(document) // 2)
    return {text for text, count in candidates.items() if count >= threshold}


def _looks_like_table(value: str) -> bool:
    lines = [line for line in value.splitlines() if line.strip()]
    if len(lines) < 2:
        return False
    pipe_lines = sum(1 for line in lines if "|" in line)
    spaced_lines = sum(1 for line in lines if re.search(r"\S\s{2,}\S", line))
    return pipe_lines >= 2 or spaced_lines >= 2


def _looks_like_formula(value: str) -> bool:
    compact = " ".join(value.split())
    if len(compact) > 300:
        return False
    math_symbols = ["=", "+", "-", "*", "/", "\\", "∑", "∫", "√", "≤", "≥", "≈", "α", "β", "γ"]
    has_symbol = any(symbol in compact for symbol in math_symbols)
    has_variable = bool(re.search(r"\b[a-zA-Z]\b", compact))
    has_number = bool(re.search(r"\d", compact))
    return has_symbol and (has_variable or has_number)


def _is_reference_section(title: str) -> bool:
    return title.strip().lower() in {"references", "reference", "bibliography", "参考文献"}


def _split_reference_entries(text: str) -> list[str]:
    value = text.strip()
    if not value:
        return []
    parts = re.split(r"\n(?=(?:\[\d+\]|\d+\.|\d+\s+)\s*)", value)
    return [part.strip() for part in parts if part.strip()]


def _clean_block_text(text: str) -> str:
    return "\n".join(line.strip() for line in text.splitlines() if line.strip())


def _split_paragraphs(content: str) -> list[str]:
    paragraphs = re.split(r"\n\s*\n", content)
    return [paragraph.strip() for paragraph in paragraphs if paragraph.strip()]


def _is_markdown_heading(line: str) -> bool:
    return bool(re.match(r"^#{1,6}\s+\S+", line.strip()))


def _parse_markdown_heading(line: str) -> tuple[int, str]:
    match = re.match(r"^(#{1,6})\s+(.+)$", line.strip())
    if not match:
        return 1, line.strip()
    return len(match.group(1)), match.group(2).strip()


def _strip_markdown_heading(text: str) -> str:
    return re.sub(r"^#{1,6}\s+", "", text).strip()


def _is_markdown_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def _is_markdown_image(line: str) -> bool:
    return bool(re.match(r"^\s*!\[[^\]]*\]\([^)]+\)\s*$", line))


def _is_list_line(line: str) -> bool:
    return bool(re.match(r"^\s*(?:[-*+]|\d+\.)\s+\S+", line))


def _heading_level(style_name: str) -> int:
    match = re.search(r"heading\s+(\d+)", style_name)
    return int(match.group(1)) if match else 1


def _guess_title(elements: list[dict]) -> str:
    for element in elements:
        if element.get("type") == "title":
            return element.get("text", "")
    for element in elements:
        text = element.get("text", "")
        if text:
            return text.splitlines()[0][:180]
    return ""


def _guess_abstract(elements: list[dict]) -> str:
    for index, element in enumerate(elements):
        if element.get("type") == "title" and element.get("text", "").lower() in {"abstract", "摘要"}:
            next_element = elements[index + 1] if index + 1 < len(elements) else {}
            return next_element.get("text", "")
    return ""


def _elements_to_pages(elements: list[dict]) -> list[dict]:
    grouped: dict[int, list[str]] = {}
    for element in elements:
        page = element.get("page_start")
        if page is None:
            continue
        grouped.setdefault(page, []).append(element.get("text", ""))

    return [
        {
            "page_number": page,
            "text": "\n\n".join(text for text in texts if text),
        }
        for page, texts in sorted(grouped.items())
    ]


def _elements_to_sections(elements: list[dict]) -> list[dict]:
    sections = []
    current_title = "Document"
    current_content: list[str] = []
    page_start = None
    page_end = None

    def flush() -> None:
        nonlocal current_content, page_start, page_end
        if not current_content:
            return
        sections.append(
            {
                "title": current_title,
                "content": "\n\n".join(current_content),
                "page_start": page_start,
                "page_end": page_end,
            }
        )
        current_content = []
        page_start = None
        page_end = None

    for element in elements:
        if element.get("type") == "title":
            flush()
            current_title = element.get("text", "") or "Document"
            continue

        text = element.get("text", "")
        if not text:
            continue
        current_content.append(text)
        page_start = page_start or element.get("page_start")
        page_end = element.get("page_end") or page_end

    flush()
    return sections
